import time
from bs4 import BeautifulSoup as bs
import undetected_chromedriver as uc
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver import DesiredCapabilities
from selenium.webdriver.common.action_chains import ActionChains 
from selenium_stealth import stealth
import os
import requests
import json
from docx import Document
options = uc.ChromeOptions() 
# options.add_argument('--headless')    #doesn't work in headless mode

# options.add_argument(f'--user-data-dir=C:\Users\{os.getlogin()}\AppData\Local\Google\Chrome\User Data\Profile 5')
# options.add_argument("--disable-extensions")
# options.add_argument('--disable-application-cache')
# options.add_argument('--disable-gpu')
# options.add_argument("--no-sandbox")
# options.add_argument("--disable-setuid-sandbox")
# options.add_argument("--disable-dev-shm-usage")
# stre = f"C:/Users/{os.getlogin()}/AppData/Local/Google/Chrome/User Data/Profile 5"
# driver = uc.Chrome(user_data_dir=stre,options=options)
options.add_argument('--disable-blink-features=AutomationControlled')

driver = uc.Chrome(options=options)
driver.get('https://babelnovel.com/')
time.sleep(5)

def SignIn():
    driver.maximize_window()
    actions = ActionChains(driver)
    more_button = driver.find_element(By.CSS_SELECTOR, " #top-bar .top-bar_user-icon__3p83F")
    actions.move_to_element(more_button).perform()
    time.sleep(2)
    login_button = driver.find_element(By.XPATH,"//button/div[contains(text(),'Log in')]")
    login_button.click()
    time.sleep(4)
    iframe = driver.find_element(By.CSS_SELECTOR,'div iframe')
    driver.switch_to.frame(iframe)
    email_field = driver.find_element(By.XPATH,"//input[@placeholder='Email']")
    password_field = driver.find_element(By.XPATH,"//input[@placeholder='Password']") 
    Login = driver.find_element(By.XPATH,"//button/p[contains(text(),'Log in')]")
    email_field.send_keys('mathurkartik93@gmail.com')
    password_field.send_keys('Pocketfm123')
    time.sleep(2)
    Login.click()
    time.sleep(8)
    # driver.switch_to.default_content()


def crawl_babel(title):
    page_start = 866    #change here
    start = page_start
    page_end = 1000    #change here
    # lis = [ 1626,1629,1635,1642,1645,1649,1666,1679,1680,1687,1688,1700]
    book_title = title
    
    document = Document()
    url=f"https://babelnovel.com/books/{book_title}/chapters/c{start}"
    driver.get(url)
    time.sleep(6)
    soup=bs(driver.page_source,"lxml")
    c=soup.find_all("p",class_="chapter_break__2SEQX")
    h=soup.find("h3",class_="chapter_title__3Dp-H")
    document.add_heading(h.text)
    for p in c:
        document.add_paragraph(p.text)
    document.add_page_break()

    time.sleep(1)
    print(f"{start} done")
    start+=1

    while start<=page_end:
    # for start in lis:
        # if (start-1)%50==0:
        #     document.save(f"{book_title}_{page_start}_{start-1}.docx")
        #     document=Document()
        #     page_start=start
        url=f"https://babelnovel.com/books/{book_title}/chapters/c{start}"
        driver.get(url)
        # time.sleep(15)
        time.sleep(4)
        driver.execute_script("window.scrollTo(0, window.scrollY + 600)")
        time.sleep(4)
        soup=bs(driver.page_source,"lxml")
        c=soup.find_all("p",class_="chapter_break__2SEQX")
        h=soup.find("h3",class_="chapter_title__3Dp-H")
        document.add_heading(h.text)
        for p in c:
            # print(p.text)
            document.add_paragraph(p.text)

        document.add_page_break()
        time.sleep(4)
        print(f"{start} done")
        start+=1

    document.save(f"{book_title}_{page_start}_{page_end}.docx")
    driver.close()
    # document.save(f"out/{book_title}.docx")


title=["her-sniper"]   #change here

SignIn()
for t in title:
    print(t)
    crawl_babel(t)