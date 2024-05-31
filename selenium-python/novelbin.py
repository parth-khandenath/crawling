import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
import time
from docx import Document


name = 'Aplhas rejected mate returns as queen'  #change here
start_chapter = 201  #change here
end_chapter = 656  #change here
count = 0
doc=Document()
file_count = 0
ll = "https://novelbin.me/book/alphas-rejected-mate-returns-as-queen"  #change here #book link
options = uc.ChromeOptions()
options.page_load_strategy = 'eager'

# option.add_experimental_option( "prefs",{'profile.managed_default_content_settings.javascript': 2})
driver=uc.Chrome(options=options)
driver.get(ll)
time.sleep(20)
soup=bs(driver.page_source,"lxml")
link = 'https://binnovel.com/novel-book/alphas-rejected-mate-returns-as-queen/cchapter-1' #change here  #chapter1 link
# link = soup.select_one('.btn-read-now').get('href')
# link = "https://novelusb.com/novel-book/shadow-slave-novel/chapter-922-little-by-little"
while count!=end_chapter:
    if count>=start_chapter-1:
        driver.get(link)
        time.sleep(3)
        soup=bs(driver.page_source,"lxml")
        link = soup.select_one('#next_chap').get('href')
        title = soup.select_one('h4')
        if title is None:
            title = soup.select_one('.chr-text').text
        else:
            title = title.text
        paragraphs = soup.select('p')
        doc.add_heading(f"{title}")
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph.text)
        doc.add_page_break()
        count+=1
        print(f"{count} done ")
        if (count - start_chapter + 1)%100 == 0:
            # file_name = './out/'+name+str(file_count)+ '.docx' 
            file_name = name+str(file_count)+ '.docx' 
            file_count+=1
            doc.save(file_name)
            doc = Document()
        continue
    driver.get(link)
    # time.sleep(7)
    soup=bs(driver.page_source,"lxml")
    link = soup.select_one('#next_chap').get('href')
    count+=1
    print(f"{count} skipped ") 

# file_name = './out/'+name + '.docx' 
file_name = name + '.docx' 
doc.save(file_name)