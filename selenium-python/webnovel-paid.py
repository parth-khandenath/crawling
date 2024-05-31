import undetected_chromedriver as uc
from selenium import webdriver
import time
from selenium.webdriver.common.by import By
from selenium.webdriver import ChromeOptions
import docx
from bs4 import BeautifulSoup as bs

def login(driver):
    driver.get("https://www.webnovel.com")
    # time.sleep(3)  # Wait for the page to load

    # Click the login button
    # login_button = self.driver.find_element(By.CSS_SELECTOR, '.login-btn')
    # login_button.click()
    # time.sleep(4)


    print("############")
    print("############")
    print("LOGIN !!!")
    print("############")
    print("############")
    time.sleep(30)          #manual login
    print('30 secs left..')
    time.sleep(30)

    # Wait for the email input to be clickable
    # email_input = WebDriverWait(self.driver, 10).until(
    #     EC.element_to_be_clickable((By.XPATH, '//input[@name="email"]'))
    # )
    # email_input.send_keys("sagarika.swain@pocketfm.com")

    # # Enter the password
    # password_input = self.driver.find_element(By.XPATH, '//input[@name="password"]')
    # password_input.send_keys("Swain@4422")

    # # Submit the login form
    # password_input.send_keys(Keys.RETURN)

    # # Wait for the user to be logged in
    # WebDriverWait(self.driver, 10).until(
    #     EC.presence_of_element_located((By.XPATH, '//span[contains(text(),"Library")]'))
    # )

paras_to_be_skipped=['Weekly Power Status',
'You may also Like',
'Background    default     yellow     dark',
'Font   Nunito Sans   Merriweather',
'Size',
'Paragraph comments',
'The total score 0.0',
'Advertisement Pornographic Personal attack Other',
'Paragraph comment',
'Paragraph comment feature is now on the Web! Move mouse over any paragraph and click the icon to add your comment.',
'Also, you can always turn it off/on in Settings.',
]
# PS C:\Program Files\Google\Chrome\Application> .\chrome.exe --remote-debugging-port=8989 --user-data-dir="C:\parth_webnovel_data"
# chrome_user_data_dir = 'C:\\Users\\Admin\\AppData\\Local\\Google\\Chrome\\User Data'
# profile = 'Profile 6'
options=ChromeOptions()
options.add_experimental_option("debuggerAddress", "localhost:8989")
options.add_argument(f'user-agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36"')
# options.add_argument(f"user-data-dir={chrome_user_data_dir}")
# options.add_argument(f"profile-directory={profile}")
driver=webdriver.Chrome(options=options)
doc=docx.Document()

start=411
end=420
first_chp_link="https://www.webnovel.com/book/ceo's-runaway-mate-and-her-cubs_21389064705112305/naughty-bunny_62426586837730674"
i=411  #chapter no of above link
docname="webnovel-allure-of-night"

# login(driver)
driver.get("https://www.webnovel.com") #/book/ceo's-runaway-mate-and-her-cubs_21389064705112305") #booklink
print("############")
print('open webnovel start chapter in 90 secs.. (new tab,only tab)')
time.sleep(60)
print("############")
print('30 secs left')
time.sleep(30)
# for _ in range(6):
#     driver.execute_script("window.scrollBy(0, 500);")
# driver.get(first_chp_link)

paras_till_now=0
while i<=end:
    time.sleep(1)
    soup=bs(driver.page_source,'lxml')
    print(soup)
    headings=soup.select('h1.dib')
    paras=soup.select('div p')
    # print(headings)
    # print(paras)
    if i>=start and i<=end:
        title=headings[i-1].text.strip()
        doc.add_heading(title)
        for p in paras[paras_till_now:]:
            if not p.text.strip() in paras_to_be_skipped:
                doc.add_paragraph(p.text.strip())
        doc.add_page_break()

        if i==end:
            doc.save(f'{docname}-{start}-{end}.docx')
        elif (i-start+1)%100==0:
            doc.save(f'{docname}-{start}-{i}.docx')
            start=i+1

    paras_till_now=len(paras)
    while(i==len(headings) and i<end):
        print('scrolling for chapter ',i+1)
        driver.execute_script("window.scrollBy(0, 500);")
        soup=bs(driver.page_source,'lxml')
        headings=soup.select('h1.dib')
    i+=1

# driver.close()
# driver.quit()