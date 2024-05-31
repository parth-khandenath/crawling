from selenium import webdriver
import undetected_chromedriver as uc
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup
import docx
import time
import json
import urllib.parse

class WebnovelScraper:
    def __init__(self, book_url, start, end, output_dict):
        self.book_url = urllib.parse.unquote(book_url)
        self.start = start
        self.end = end
        self.output_dict = output_dict
        self.counter = 36
        self.doc = docx.Document()
        self.driver = None

    def setup_driver(self):
        chrome_options = uc.ChromeOptions()
        # chrome_options.add_argument('--headless')
        chrome_options.add_argument('--disable-gpu')
        # chrome_options.add_argument('--_csrfToken=2d145ea1-113c-4967-9ab8-d4ea1242587a; webnovel_uuid=1716377831_552466962; webnovel-content-language=en; webnovel-language=en; bookCitysex=1; _fbp=fb.1.1716377835877.1197789475; _gid=GA1.2.2101934915.1716377836; dontneedgoogleonetap=1; e2=%7B%22pid%22%3A%22bookstore%22%2C%22l1%22%3A%221%22%7D; e1=%7B%22pid%22%3A%22bookstore%22%2C%22l1%22%3A%2299%22%7D; _ga_GG9P0S9KJK=GS1.1.1716381965.2.0.1716381965.60.0.0; AMP_TOKEN=%24NOT_FOUND; uid=4326050085; ukey=uOFEhE9dzVu; checkInTip=1; __gads=ID=1d6141d3632d07eb:T=1716377838:RT=1716384032:S=ALNI_MaT37IQoLdXefD6Ip43C7LGDFbvsA; __gpi=UID=00000e27edd041ee:T=1716377838:RT=1716384032:S=ALNI_MZtZN-YWEBS4dS9pwPMee7g4Q4gIA; __eoi=ID=e6ac7a7a67c2ebff:T=1716377838:RT=1716384032:S=AA-Afja_jSWCRSOTLdnRqGlopKI1; _ga=GA1.1.1953967122.1716377836; _ga_PH6KHFG28N=GS1.1.1716381970.2.1.1716384192.60.0.0')
        # chrome_service = Service('path/to/chromedriver')  # Adjust the path to the chromedriver
        self.driver = uc.Chrome(options=chrome_options)

    def login(self):
        self.driver.get("https://www.webnovel.com")
        time.sleep(3)  # Wait for the page to load

        # Click the login button
        # login_button = self.driver.find_element(By.CSS_SELECTOR, '.login-btn')
        # login_button.click()
        # time.sleep(4)


        print("############")
        print("############")
        print("LOGIN !!!")
        print("############")
        print("############")
        time.sleep(60)          #manual login
        print('1 min left..')
        time.sleep(30)

        # Wait for the email input to be clickable
        # email_input = WebDriverWait(self.driver, 10).until(
        #     EC.element_to_be_clickable((By.XPATH, '//input[@name="email"]'))
        # )
        # email_input.send_keys("sagarika.swain@pocketfm.com")

        # # Enter the password
        # password_input = self.driver.find_element(By.XPATH, '//input[@name="password"]')
        # password_input.send_keys("Swain@4422")

        # # Submit the login form
        # password_input.send_keys(Keys.RETURN)

        # # Wait for the user to be logged in
        # WebDriverWait(self.driver, 10).until(
        #     EC.presence_of_element_located((By.XPATH, '//span[contains(text(),"Library")]'))
        # )

    def scrape(self):
        self.setup_driver()
        self.login()
        # self.driver.get(self.book_url)
        # time.sleep(3)  # Wait for the page to load

        # soup = BeautifulSoup(self.driver.page_source, 'html.parser')
        # first_chap_url = soup.select_one('a#j_read')['href']
        # chapter_1_link = 'https://www.webnovel.com' + first_chap_url

        # chap = chapter_1_link.split('/')[-1]
        # chapter_id = chap.split('_')[-1]
        book_id='23138127806346305'
        chapter_id = '62545417056862272'
        # book = chapter_1_link.split('/')[-2]
        # book_id = book.split('_')[-1]
        # title = book.split('_')[0]

        # cookies = self.driver.get_cookies()
        csrf_token = '2d145ea1-113c-4967-9ab8-d4ea1242587a'
        # for cookie in cookies:
        #     if cookie['name'] == '_csrfToken':
        #         csrf_token = cookie['value']
        #         break

        self.parse_chapters(book_id, chapter_id, csrf_token)
        # self.driver.quit()

    def parse_chapters(self, book_id, chapter_id, csrf_token):
        url = f"https://www.webnovel.com/go/pcm/chapter/getContent?_csrfToken={csrf_token}&_fsae=0&bookId={book_id}&chapterId={chapter_id}&encryptType=3&font=Merriweather&_=1716384049110"
        self.driver.get(url)
        time.sleep(10)  # Wait for the content to load

        body_json = self.driver.find_element(By.TAG_NAME, 'body').text
        data = json.loads(body_json)['data']['chapterInfo']
        vip = data['vipStatus']

        # if vip == 0:
        contents = data['contents']
        chapter_name = data['chapterName']
        chapter_id = data['chapterId']

        if self.counter >= self.start:
            self.doc.add_heading(str(self.counter), 0)
            self.doc.add_heading(chapter_name, 0)

            for content in contents:
                text = content['content'].strip('<p>').strip('</p>')
                self.doc.add_paragraph(text)
            self.doc.add_page_break()

        if self.counter == self.end:
            # self.output_dict['output'] = self.doc
            self.doc.save(f'webnovel_content-{self.start}-{self.end}.docx')
            return

        if (self.counter-self.start+1)%100==0:
            self.doc.save(f'webnovel_content-{self.start}-{self.counter}.docx')
            self.start=self.counter+1
            self.doc=docx.Document()
        self.counter += 1
        next_chapter_id = data['nextChapterId']
        self.parse_chapters(book_id, next_chapter_id, csrf_token)
        # else:
            # self.output_dict['output'] = self.doc

if __name__ == "__main__":
    output_dict = {}
    start=36
    end=400
    scraper = WebnovelScraper(book_url="https://www.webnovel.com/book/the-lycan's-sin_23138127806346305", start=start, end=end, output_dict=output_dict)
    scraper.scrape()
    # doc = output_dict['output']
    # doc.save('webnovel_content-{start}-{end}-fin.docx')
