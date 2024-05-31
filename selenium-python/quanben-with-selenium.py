import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs
import pandas as pd
import time
import docx
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
import math
from selenium import webdriver
import re 

def read_number(str):
    pattern = r'第(.*?)章'
    lst = re.findall(pattern, str)
    try:
        return int(lst[0])
    except:
        return -1

head = {
    "Title" :[],
    "URL":[],
    "book id":[],
    "Author":[],
    "Description":[],
    "Status":[],
    "Views":[],
    "Likes":[],
    "Tags":[],
    # "Chapters":[],
    "Genre":[]
}

doc = docx.Document()

options = uc.ChromeOptions() 
options.add_argument('--disable-blink-features=AutomationControlled')
# options.add_argument('--headless')
options.add_argument(f"--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.159 Safari/537.36")

driver = uc.Chrome(options=options)

book_chinese_name='guiyudangan'     #change here
start_chapter=1601                     #change here
end_chapter=1895                    #change here
site=f'https://www.quanben-xiaoshuo.com/n/{book_chinese_name}/xiaoshuo.html'

driver.get(site)
chpt_title_eles=driver.find_elements(By.CSS_SELECTOR,'.list a span')
chpt_titles=[]
for ele in chpt_title_eles:
    chpt_titles.append(ele.text)
last_chapter=read_number(chpt_titles[-1])
chpt_link_eles=driver.find_elements(By.CSS_SELECTOR,'.list a')
chpt_links=[]
for ele in chpt_link_eles:
    chpt_links.append(ele.get_attribute('href'))
# if(end_chapter>last_chapter): #some/all of the requested chapters are unavailable
#     if(start_chapter>last_chapter): #all requested chapters are unavailable
#         doc.add_paragraph("All the requested chapters are unavailable.")
#     else: #fetching few initial available chapters 
#         doc.add_paragraph("Some of the requested chapters maybe unavailable.")
#         end_chapter=last_chapter #update end chapter to last available chapter
i=0
while i<(len(chpt_links)):
    # ch_no=read_number(chpt_titles[i])
    # if ch_no>=start_chapter and ch_no<=end_chapter:
    if i+1<start_chapter:
        i+=1
        continue
    print(i+1)
    driver.get(chpt_links[i])
    chpt_title=driver.find_element(By.CSS_SELECTOR,'.title').text
    doc.add_heading(chpt_title)
    chpt_content=driver.find_elements(By.CSS_SELECTOR,'p')
    for para_ele in chpt_content:
        para=para_ele.text
        para=para.strip()
        if para=='':
            continue
        doc.add_paragraph(para)
    doc.add_page_break()
    if (i+1)%100==0:
        doc.save(f'quanben-{book_chinese_name}-{start_chapter}-{i+1}.docx')
        doc = docx.Document()
        start_chapter=i+1+1
    i+=1

doc.save(f'quanben-{book_chinese_name}-{start_chapter}-{end_chapter}.docx')