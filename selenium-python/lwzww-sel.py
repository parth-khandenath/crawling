import time
import re
import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from docx import Document

class lwzww_Scraper:
    def __init__(self):
        self.start_chapter = 1 # change here
        self.end_chapter = 20 # change here
        self.book_name = "Death heartbeat" # change here
        self.start_page = (self.start_chapter-1) // 100 + 1
        self.end_page = (self.end_chapter-1) // 100 + 1
        self.start_urls = [f'http://www.lwzww.com/book/5ilf/{str(i)}/' for i in range(self.start_page, self.end_page + 1)]
        self.doc = Document()

        headers = {
            'accept': 'application/json, text/plain, */*',
            'accept-language': 'en-GB,en-US;q=0.9,en;q=0.8',
            'content-type': 'application/json',
            'cookie': 'qq_domain_video_guid_verify=7f3ccf505389bdc2; _qimei_uuid42=183010e113a1002f9862363182dec709edd78e889d; _qimei_fingerprint=1c609ff4531c6a835b185cfa96d323f7; pgv_pvid=2097815516; video_platform=2; video_guid=7f3ccf505389bdc2; _qimei_q36=; _qimei_h38=2db878989862363182dec70902000001218301; pgv_info=ssid=s3417771617',
            'origin': 'http://www.lwzww.com',
            'referer': 'http://www.lwzww.com/',
            'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Google Chrome";v="122"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"Windows"',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-site',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
        }

        options = uc.ChromeOptions()
        for key, value in headers.items():
            options.add_argument(f"--{key}={value}")
        self.driver = uc.Chrome(options=options)

    def scrape(self):
        for url in self.start_urls:
            self.driver.get(url)
            time.sleep(10)  # Adjust time as needed for page to load
            a_tags = self.driver.find_elements(By.CSS_SELECTOR, '.chapter-5c993fac a')
            chpt_links = []
            for a in a_tags:
                onclick_attr = a.get_attribute('onclick')
                match = re.search(r"location\.href='(.*?)'", onclick_attr)
                if match:
                    url = match.group(1)
                    chpt_links.append('http://www.lwzww.com' + url)
            c=0
            for link in chpt_links:
                c+=1
                if c>=self.start_chapter and c<=self.end_chapter:
                    self.parse_chapter(link)

        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')

    def parse_chapter(self, url):
        self.driver.get(url)
        time.sleep(2)  # Adjust time as needed for page to load
        title = self.driver.find_element(By.CSS_SELECTOR, '.h1-5c993fac').text
        paras = self.driver.find_elements(By.CSS_SELECTOR, '#chaptercontent p')
        self.doc.add_heading(title)
        for para in paras:
            para_text = para.text.strip()
            if para_text:
                self.doc.add_paragraph(para_text)
                self.doc.add_paragraph('')
        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')

    def close_driver(self):
        self.driver.quit()

if __name__ == "__main__":
    scraper = lwzww_Scraper()
    scraper.scrape()
    scraper.close_driver()
