import docx, time
from  selenium import webdriver
import undetected_chromedriver as uc
from bs4 import BeautifulSoup as bs



book_url = "https://www.17k.com/book/2820131.html"    #change here
book_name = "Qi Men Ming Shu"     #change here
start = 1     #change here
end = 234      #change here
count = start
doc = docx.Document()

options = uc.ChromeOptions()
# options.add_argument('--headless')
options.page_load_strategy = 'eager'
driver = uc.Chrome(options=options)

driver.get(book_url)
time.sleep(10)
soup = bs(driver.page_source,"lxml")

if '/book/' in driver.current_url:
    chapter_list = soup.select_one('dt.read a').get('href') 
    chapter_list = 'https://www.17k.com'+chapter_list
    driver.get(chapter_list)
    time.sleep(5)
    soup = bs(driver.page_source,"lxml")
if '/list/' in driver.current_url:
    first_chapter_url = soup.select_one('dl.Volume dd a').get('href')
    first_chapter_url = 'https://www.17k.com'+first_chapter_url

driver.get(first_chapter_url)
time.sleep(5)
soup = bs(driver.page_source,"lxml")

while count<end:

    title = soup.select_one('div.readAreaBox.content h1').text
    content = soup.select('div.readAreaBox.content p')

    if soup.select_one('div.free'):
        break
    else:
        if count>=start:
            doc.add_heading(title)
            for c in content:
                if c:
                    doc.add_paragraph(c.text.strip())
            doc.add_page_break()
            doc.save(f'{book_name}.docx')
        
        next_chapter_url = soup.select_one('.NextPrevBtn .nextChapter').get('href')
        print(next_chapter_url)
        next_chapter_url = 'https://www.17k.com' + next_chapter_url
        driver.get(next_chapter_url)
        time.sleep(3)
        soup = bs(driver.page_source,"lxml")

    count+=1

doc.save(f'{book_name}.docx')