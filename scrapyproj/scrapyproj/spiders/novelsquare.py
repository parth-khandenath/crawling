import scrapy
import os
from docx import Document
import docx

class NovelsquareSpider(scrapy.Spider):
    name = "novelsquare"
    allowed_domains = ["www.novelsquare.com","www.welikereading.com"]
    start_urls = ["https://www.novelsquare.com/chapter/chapter-209280-drluna--67da8177/"]      #put url of start chapter #change here
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        # 'DOWNLOAD_DELAY':1
    }
    def __init__(self):
        self.doc = docx.Document()
        self.start = 1  #change here
        self.end =500   #change here
        self.count = self.start
        self.bookname = "Dr Luna"  #change here
        self.doc = docx.Document()
        self.priority = 100000000
        self.error_end = False

    def parse(self, response):
        heading=response.css(".my-8::text").get()
        content=response.css(".prose")
        all_text_content = content.css('::text').getall()
        
        #Add data to file
        self.doc.add_heading(heading, level=1)
        for text in all_text_content:
            text=text.strip()
            if text=='':
                continue
            self.doc.add_paragraph(text)
        self.doc.add_page_break()

        if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
                    self.doc.save(f'{self.bookname}-{((self.count-2)//100)*100 +1}-{self.count}.docx')
                    self.doc = docx.Document()

        self.count += 1

        if self.count <= self.end:
            relative_url= response.css('.decoration-solid ::attr(href)').getall()[1]
            if self.count<50:
                next_chapter_url= "https://www.novelsquare.com" + relative_url
            else:
                next_chapter_url=   relative_url
            next_chapter_url=next_chapter_url.replace('welikereading','novelsquare')
            print("NEXT CHAPTER AT ", next_chapter_url)
            yield response.follow(next_chapter_url, callback=self.parse)
        else:
                print("ELSE MAI HAI HUM")
                self.doc.save(f'{self.bookname}-{self.count}.docx')
                self.error_end = True
    
    def closed(self,reason):
        print(self.count)

        print(reason)
        print("CLOSEDDDDDDDDDDDDDDDD")
        self.doc.save(f'{self.bookname}.docx') 