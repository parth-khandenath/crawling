import scrapy
import docx

class seventeenkContent(scrapy.Spider):
    name = 'seventeenkContent'
    allowed_domains = ['www.xianqihaotianmi.org']
    def __init__(self, **kwargs):
        self.start_urls = ['http://www.xianqihaotianmi.org/book/42864.html']
        self.start_chapter = 101 #change here
        self.end_chapter = 500 #change here
        self.book_name = 'spy war elite' #change here
        self.doc = docx.Document()
        self.priority = 1000000
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response):
        # all_anchors=response.css('.list-charts a').getall()
        all_links=response.css('.list-charts a::attr("href")').getall()
        all_texts=response.css('.list-charts a::text').getall()
        for idx in range(len(all_texts)):
            if '章' in all_texts[idx]: #it is a chapter
                if idx+1>=self.start_chapter and idx+1<=self.end_chapter: #within range
                    self.priority-=1
                    yield response.follow(url='http://www.xianqihaotianmi.org'+all_links[idx],callback=self.parse_chapter,priority=self.priority)
                elif idx+1>self.end_chapter:
                    break

    def parse_chapter(self,response):
        chp_title=response.css('.panel-heading::text').get()
        chpt_paras=response.css('.content-body::text').getall()
        # print(chp_title)
        # print(chpt_paras)
        self.doc.add_heading(chp_title)
        for para in chpt_paras:
            para=para.strip()
            if para=='':
                continue
            self.doc.add_paragraph(para)
            self.doc.add_paragraph('')
        self.doc.add_page_break()

    def closed(self,reason):
        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')
    