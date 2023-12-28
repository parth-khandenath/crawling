import scrapy
import docx
import requests

class PowertoolContentZonghengSpider(scrapy.Spider):
    name = 'zongheng'
    allowed_domains = ['book.zongheng.com','read.zongheng.com']
    def __init__(self, **kwargs):
        self.start_urls = ['https://huayu.zongheng.com/showchapter/1119895.html']
        self.start_chapter = 1
        self.end_chapter = 2
        # self.output_dict = kwargs['output_dict']
        self.doc = docx.Document()
        self.priority = 1000000
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response):
        url = response.url.split('/')
        if 'detail' in url:  #for https://www.zongheng.com/detail/777173 type of links
            if url[-1]=='':  #when link has / at end
                book_id = url[-2]
            else:            #when link doesn't have / at end
                book_id = url[-1]
            url = "https://bookapi.zongheng.com/api/chapter/getChapterList"
            payload = f'bookId={book_id}'
            headers = {
            'Accept': 'application/json, text/plain, */*',
            'Accept-Language': 'en-GB,en-US;q=0.9,en;q=0.8',
            'Connection': 'keep-alive',
            'Content-Type': 'application/x-www-form-urlencoded',
            'Cookie': 'zhffr=0; ZHID=E0B05FA646E8203EBE35E919B84D979D; zh_visitTime=1699290705793; Hm_lvt_c202865d524849216eea846069349eb9=1699290790; v_user=https%3A%2F%2Fbook.zongheng.com%2F%7Chttps%3A%2F%2Fhuayu.zongheng.com%2Fbook%2F1292301.html%7C96503595; ver=2018; PassportCaptchaId=08d777141f3827f6ec5374bf63e599a5; Hm_lpvt_c202865d524849216eea846069349eb9=1701255318; sensorsdata2015jssdkcross=%7B%22distinct_id%22%3A%2218ba36c5459515-0749f66ac4ed53-16525634-1187056-18ba36c545a152c%22%2C%22%24device_id%22%3A%2218ba36c5459515-0749f66ac4ed53-16525634-1187056-18ba36c545a152c%22%2C%22props%22%3A%7B%22%24latest_traffic_source_type%22%3A%22%E7%9B%B4%E6%8E%A5%E6%B5%81%E9%87%8F%22%2C%22%24latest_referrer%22%3A%22%22%2C%22%24latest_referrer_host%22%3A%22%22%2C%22%24latest_search_keyword%22%3A%22%E6%9C%AA%E5%8F%96%E5%88%B0%E5%80%BC_%E7%9B%B4%E6%8E%A5%E6%89%93%E5%BC%80%22%7D%7D',
            'Origin': 'https://www.zongheng.com',
            'Referer': 'https://www.zongheng.com/',
            'Sec-Fetch-Dest': 'empty',
            'Sec-Fetch-Mode': 'cors',
            'Sec-Fetch-Site': 'same-site',
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
            'sec-ch-ua': '"Google Chrome";v="119", "Chromium";v="119", "Not?A_Brand";v="24"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"macOS"'
            }

            chapters_response = requests.request("POST", url, headers=headers, data=payload)
            if chapters_response.status_code==200:
                response_json=chapters_response.json()
                chapter_list=(((response_json['result']['chapterList'])[1])['chapterViewList']) #all chapters
                chapters_in_range=[] #free chapters within requested range

                for indx in range(self.start_chapter-1, min(self.end_chapter,len(chapter_list))):
                    current_chapter=chapter_list[indx]
                    if current_chapter['price']>0: #VIP/paid chapters start
                        break
                    else:   #free chapters
                        chapters_in_range.append(current_chapter)

                for chapter in chapters_in_range:
                    chapter_id=chapter['chapterId']
                    chapter_link=f"https://read.zongheng.com/chapter/{book_id}/{chapter_id}.html"
                    yield response.follow(url=chapter_link,callback=self.forcontent,priority=self.priority)
                    self.priority-=1

            else:
                print("Could not fetch chapters...")
                print(f"Status code:{chapters_response.status_code}") 
            
        else: #for https://huayu.zongheng.com/showchapter/1262587.html type of links
            if url[3] == 'showchapter':    
                chapter_url = response.url
            else:
                chapter_url = response.css('.all-catalog::attr("href")').get()
            yield response.follow(url=chapter_url,callback=self.chapters,dont_filter=True)
    
    def chapters(self,response):
        all_links = response.css('.col-4 a::attr("href")').getall()
        vip_links = response.css('.vip a::attr("href")').getall()
        all_free_links = [link for link in all_links if link not in vip_links]
        last_free_chapter=len(all_free_links)

        if(self.end_chapter>last_free_chapter): #some/all of the requested chapters are VIP
            if(self.start_chapter>last_free_chapter): #all requested chapters are VIP
                self.doc.add_paragraph("All the requested chapters are VIP/paid chapters.")
                return
            else: #fetching few initial free chapters 
                self.doc.add_paragraph("Some of the requested chapters are VIP/paid chapters.")
                self.end_chapter=last_free_chapter #update end chapter to last free chapter

        for i in range(self.start_chapter,min(self.end_chapter+1,len(all_free_links)+1)):
            current_chapter_link = all_free_links[i-1]
            yield response.follow(url=current_chapter_link,callback=self.forcontent,priority=self.priority)
            self.priority-=1    

    def forcontent(self,response):
        title = response.css('.title_txtbox::text').get()
        content = response.css('.content p::text').getall()
        self.doc.add_heading(title)
        for paragraph in content:
            self.doc.add_paragraph(paragraph)
        self.doc.add_page_break()
        # self.output_dict['output'] = self.doc 

    def closed(self,reason):
        self.doc.save('trial.docx')  
