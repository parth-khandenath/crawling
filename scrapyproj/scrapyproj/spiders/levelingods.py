import scrapy
import docx

class levelingods(scrapy.Spider):
    name = 'levelingods'
    def __init__(self, **kwargs):

        self.name = 'IBAWWMB Novel'   #change this   
        self.book_url = 'https://www.levelingods.shop/p/isnt-being-wicked-woman-much-better.html'   #change this
        
        self.start_urls = [ self.book_url ]     
        self.start = 1  # like this
        self.end = 167
        # self.output_dict = {}
        self.doc = docx.Document()
        self.priority = 100000000
        self.website = ''
        self.book = ''
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response): #on book page
        links=response.css("#post-body a::attr(href)").getall()
        links=links[1:]
        for i in range(len(links)):
            yield response.follow(url = links[i], callback=self.parse_chapter, priority = self.priority, meta={'no':i+1})
            self.priority -= 1
            # break
    
    def parse_chapter(self,response): #on chapter page
        print('doing chapter:',response.meta['no'])
        # title = response.css('#post-body b span::text').get()
        # if not title or len(title):
        #     title = response.css('#post-body b::text').get()
        # self.doc.add_heading(title)
        # if response.meta['no']==1:
        # content = response.css('#post-body div span::text').getall()
        content = response.css('.separator *::text').getall()
        if len(content)==0:
            content = response.css('#post-body span::text').getall()
        self.doc.add_heading(content[0])
        content=content[1:]
        for para in content:
            # print(para)
            p = para.strip() #removing leading and trailing whitespaces
            if p=='':
                continue
            self.doc.add_paragraph(p)
            self.doc.add_paragraph('') #adding a blank line before a new para
        self.doc.add_page_break()
        # below part not understood
        # if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
        #     self.doc.save(f'{self.name}--{((self.count-2)//100)*100} -{self.count}.docx')
        #     self.doc = docx.Document()
        # self.count += 1

    def closed(self,reason):
        self.doc.save(f'{self.name}-{self.start}-{self.end}.docx') 
    
        