import scrapy
import docx
from bs4 import BeautifulSoup as bs

class Ixdzs(scrapy.Spider):
    name = 'ixdzs'
    def __init__(self, **kwargs):

        self.name = 'concubine Row 644'   #change this   
        self.book_url = 'https://ixdzs8.tw/read/150839/'   #change this
        
        self.start_urls = [ self.book_url ]     
        self.start = 1  # like this
        self.end = 200
        self.count = self.start
        self.output_dict = {}
        self.doc = docx.Document()
        self.priority = 100000000
        self.website = ''
        self.book = ''
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response):
        print(response.url)
        for part in response.url.split('/'):
            if 'ixdzs' in part:
                self.website = "https://"+part
                print('website = ',self.website)
                self.book =self.website
                continue
            if len(self.website) >0:
                self.book += '/'+part
                # print(self.book)
        if self.book[-1]!='/':
            self.book += '/'
        for i in range(self.start,self.end+1):
            yield response.follow(url = self.book + f"p{i}.html", callback=self.forcontent, priority = self.priority, meta = {'count':self.count})
            self.priority -= 1
            self.count +=1
    
    def forcontent(self,response):
        count=response.meta['count']
        print(count)
        title = response.css('.page-content h3 ::text').get()
        soup = bs(response.text,"lxml")
        content = soup.select('#page p')
        self.doc.add_heading(title)
        for para in content:
            p = para.text.strip()
            if p=='':
                continue
            self.doc.add_paragraph(p)
            self.doc.add_paragraph('')
        self.doc.add_page_break()
        
        if ((count) % 100 == 0 and (count - self.start) != 0) or count >= self.end:
            self.doc.save(f'{self.name}--{self.start if((count - 100)<self.start) else (count -100)} -{count}.docx')
            self.doc = docx.Document()
    
        