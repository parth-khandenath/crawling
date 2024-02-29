import scrapy
import docx
import time
import requests
# chapters 88-91    189-192
class hotnovelpub(scrapy.Spider):
    name = 'hot_novel_pub'
    allowed_domains = ['hotnovelpub.com']
    req=0

    def __init__(self, **kwargs):
        self.start_urls = []
        self.start_chapter = 1 #change here
        self.end_chapter = 200 #change here
        self.book_name = "Contract Marriage The Replacement Groom" #change here
        for i in range(self.start_chapter,self.end_chapter+1): #change below link
            self.start_urls.append(f'https://hotnovelpub.com/hotnovelpub-contract-marriage-the-replacement-groom/chapter-{str(i)}.{str(2669017+i)}')
        self.doc = docx.Document()
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response): 
        self.req+=1
        if self.req%50==0:
            time.sleep(6)
        chp_title=response.css('.chapter-name h4::text').get()
        # print(chp_title)
        chpt_paras=response.css('#content-item p::text').getall()

        url=response.url
        param=url.split('/')[-1]
        url = f"https://hotnovelpub.com/server/getContent?slug=hotnovelpub-contract-marriage-the-replacement-groom/{param}&maxWidth=640&mode=light&lineHeight=30&fontSize=20&font=Roboto"
        headers = {
        'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Google Chrome";v="122"',
        'Accept': 'application/json, text/plain, */*',
        'Referer': f'https://hotnovelpub.com/hotnovelpub-contract-marriage-the-replacement-groom/{param}',
        'sec-ch-ua-mobile': '?0',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
        'sec-ch-ua-platform': '"Windows"'
        }
        self.req+=1
        if self.req%50==0:
            time.sleep(6)
        rest_paras =[] 
        response = requests.request("GET", url, headers=headers)
        for content in response.json()['data']:
            cont=content.replace('\t',' ')
            rest_paras.extend(cont.split('\n'))

        # chpt_paras2=response.css('#content-item-0 p p::text').getall()
        # chpt_paras3=response.css('#content-item-1 p p::text').getall()
        # chpt_paras4=response.css('#content-item-2 p p::text').getall()
        # print(chpt_paras)
        # print(rest_paras)
        # print(chpt_paras2)
        # print(chpt_paras3)
        # print(chpt_paras4)
        lst=[]
        lst.extend(chpt_paras)
        lst.extend(rest_paras)
        # lst.extend(chpt_paras3)
        # lst.extend(chpt_paras4)
        self.doc.add_heading(chp_title)
        for para in lst:
            para=para.strip()
            if para=='':
                continue
            self.doc.add_paragraph(para)
            self.doc.add_paragraph('')

    def closed(self,reason):
        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')
    
# https://hotnovelpub.com/hotnovelpub-contract-marriage-the-replacement-groom/chapter-1.2669018