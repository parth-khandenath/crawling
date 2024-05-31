import scrapy
import docx
import time
from bs4 import BeautifulSoup as bs

class novelHall(scrapy.Spider):
    name = 'novelhall'
    def __init__(self, **kwargs):

        self.bookname = "The Alpha's Addiction"  #change here
        book_url = 'https://www.novelhall.com/2022The-Alpha-s-Addiction-25598/'    #change here
        self.start_urls = [ book_url ]
        self.start = 1  # like this   #change here
        self.end =150     #change here
        self.count = self.start
        self.output_dict = {}
        self.doc = docx.Document()
        self.priority = 100000000
        self.error_end = False
                
    custom_settings={
        'CONCURRENT_REQUESTS':1,
        'USER_AGENT': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36'
    }
    
    def parse(self,response):
        print("**************")
        chapter_links = response.css('#morelist a::attr(href)').getall()
        print("length", len(chapter_links))
        
        # if self.end>len(chapter_links):
        self.end=len(chapter_links)
        for link in chapter_links[self.start-1:self.end]:
            if self.error_end ==False:
                yield response.follow(url=link,callback=self.forcontent,priority = self.priority)
                self.priority -= 1
                if (self.count%100)==0:
                    time.sleep(5)
            else:
                break

    def forcontent(self,response):  
        try:
            if self.error_end == False:
                title = response.css('h1::text').get()
                content = response.css('#htmlContent *::text').getall()
                self.doc.add_heading(title)
                for c in content:
                    c=c.strip()
                    if c=='':
                        continue
                    self.doc.add_paragraph(c)
                self.doc.add_page_break()
                if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
                    self.doc.save(f'{self.bookname}-{((self.count-2)//100)*100  +1}-{self.count}.docx')
                    self.doc = docx.Document()
                self.count += 1
        except:
            if self.error_end ==False:
                self.doc.save(f'{self.bookname}-{self.count}.docx')
                self.error_end = True
            else:
                print('following error 307 encountered')

    def closed(self,reason):
        self.doc.save(f'{self.bookname}.docx') 