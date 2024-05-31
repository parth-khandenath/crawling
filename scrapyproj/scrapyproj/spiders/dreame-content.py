import scrapy
from docx import Document

class dreamecontentSpider(scrapy.Spider) :
    name = "dreamecontent"
    start_urls = ["https://www.dreame.com/story/3686211328-his-unsuspecting-mate/1121721600-chapter-1.html"]  #put 1st chapter link here
    book_name = "His Unsuspecting Mate"
    start = 1
    end = 13
    count = start
    doc = Document()

    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS':1,
        'RETRY_TIMES': 5,
        'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
    }

    def parse(self, response):
        heading = response.css("div._cid__chapter-head__5KEIg h1::text").get()
        if heading:
            self.doc.add_heading(heading, level=1)
            print("Heading is:", heading)

        content = response.css("div._cid__chapter-body__BZsTA::text").getall()
        # print("Content is  ******* :",content)
        for para in content:
            self.doc.add_paragraph(para)
            # print(para)
        self.doc.add_page_break()

        next_chapter_links = response.css("li.next a::attr(href)").getall()
        next_chapter_link = next_chapter_links[-1]
        print("Next chapter link is :", next_chapter_link)
        if self.count < self.end :
            self.count += 1
            yield response.follow(next_chapter_link, callback=self.parse)
        else :
            self.doc.save(f"{self.book_name}-{self.start} to {self.end}.docx")

    def closed(self, reason) :
        print(f"Document saved as {self.book_name}-{self.start} to {self.end}.docx")