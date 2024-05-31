
import scrapy
from docx import Document

class booktokiSpider(scrapy.Spider):
    name="booktoki_specific"
    allowed_domains = ["booktoki333.com","booktoki334.com"]
    start_urls = ["https://booktoki334.com/novel/4294005?stx=%EA%B9%80+%EB%B9%84%EC%84%9C%EA%B0%80+%EC%99%9C+%EA%B7%B8%EB%9F%B4%EA%B9%8C&book=%EC%99%84%EA%B2%B0%EC%86%8C%EC%84%A4&spage=1"]  #put start chapter link here
    book_name="Why is secretary Kim like that"
    start_chp = 1
    end_chp = 4
    start= start_chp
    chapters=4
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        # 'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'RETRY_TIMES': 5,  # Adjust the number of retries
        'DOWNLOAD_DELAY': 5,  # Add delay between requests
    }
    


    def __init__(self, *args, **kwargs):
        super(booktokiSpider, self).__init__(*args, **kwargs)
        self.doc = Document()
        self.curr_chp = self.start_chp

    def parse(self, response):

        print("chp: ",self.curr_chp)
        heading=response.css(".toon-title::text").get()
        paragraphs=response.css("#novel_content p::text").getall()

        self.doc.add_heading(heading, level=1)
        for paragraph in paragraphs:
            self.doc.add_paragraph(paragraph)
        self.doc.add_page_break()
        
        if (self.curr_chp-self.start+1)%100==0:
            self.doc.save(f'{self.book_name}-{self.start}-{self.curr_chp}.docx')
            self.start=self.curr_chp+1
            self.doc=Document()
        self.curr_chp+=1
        if self.curr_chp <= self.chapters :
            next_url=response.css(".toon-nav a#goNextBtn[href^='https://booktoki334.com/novel/']::attr(href)").get()
            if not next_url:
                next_url=response.css("a#goNextBtn::attr(href)").get()
                if not next_url or next_url=="#next":   #reachedd last chapter
                    self.doc.save(f'{self.book_name}-{self.start}-{self.curr_chp-1}.docx')   
                    return
            print(next_url)
            yield response.follow(next_url, callback=self.parse)

    def closed(self, reason):
        self.doc.save(f'{self.book_name}-{self.start}-{self.end_chp}.docx')   
