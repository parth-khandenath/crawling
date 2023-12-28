import scrapy
import docx
import re

def extract_chapt(text):
    pattern = r'第(\d+)章'
    matches = re.findall(pattern, text)
    if len(matches)==0:
        return -1
    return int(matches[0])

class quanben_xiaoshuo(scrapy.Spider):
    name = 'quanben_xiaoshuo'

    def __init__(self,name = None, **kwargs,):
        # self.book_url='https://www.quanben.io/n/wuxianzhuixiong' #put book url here
        self.book_url='https://www.quanben-xiaoshuo.com/n/wuxianzhuixiong' #alternate book url
        # self.start_urls = [ self.book_url+'/list.html'] #table of content(chapters list)
        self.start_urls = [ self.book_url+'/xiaoshuo.html'] #alternate table of content(chapters list)
        self.book_name = 'Endless Pursuit of the Culprit' #put book name here
        self.start = 1 #startchapter
        self.end = 20  #endchapter
        self.doc = docx.Document()
        self.priority = 100000000
        self.total_chapters = 0
                
    custom_settings={'DUPEFILTER_CLASS': 'scrapy.dupefilters.BaseDupeFilter',
        'CONCURRENT_REQUESTS': 1,
        'USER_AGENT':'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36'
    }
    
    def parse(self,response):
        chapts_link=response.css('.list li a::attr("href")').getall()
        chapts_titles=response.css('.list li a span::text').getall()
        self.total_chapters =extract_chapt(chapts_titles[-1])

        if(self.end>self.total_chapters): #some/all of the requested chapters are unavailable
            if(self.start>self.total_chapters): #all requested chapters are unavailable
                self.doc.add_paragraph("All the requested chapters are not released yet.")
                return
            else: #fetching few initial available chapters 
                self.doc.add_paragraph("Some of the requested chapters are not released yet.")
                self.end=self.total_chapters #update end chapter to last available chapter

        self.doc.add_paragraph("Printing available chapters out of the requested chapters.")
        for i in range(len(chapts_titles)):
            ch_no=extract_chapt(chapts_titles[i])
            print(ch_no)
            if ch_no>=self.start and ch_no<=self.end:
                yield response.follow(url=chapts_link[i],callback=self.parse_chapter,priority = self.priority)
                self.priority -= 1            
         
    def parse_chapter(self,response):
        title = response.css('h1.title ::text').get()
        content = response.css('#articlebody p ::text').getall()
        self.doc.add_heading(title)
        for para in content:
            para=para.strip()
            if para=='':
                continue
            self.doc.add_paragraph(para)
            self.doc.add_paragraph('')
        self.doc.add_page_break()
        
    def closed(self,reason):
        self.doc.save(f'{self.book_name}-{self.start}-{self.end}.docx')