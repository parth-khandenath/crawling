import scrapy
import docx

class mrfengSpider(scrapy.Spider):
    name = 'mrfeng'
    def __init__(self, **kwargs):

        self.name = 'Mr. Feng'   #change this   
        self.book_url = 'https://www.xiaoshuo.com.tw/19604/'   #change this
        
        self.start_urls = [ self.book_url ]     
        self.start = 251  # like this
        self.end = 500
        # self.output_dict = {}
        self.doc = docx.Document()
        self.priority = 100000000
        self.website = ''
        self.book = ''
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response): #on book page
        for i in range(self.start,self.end+1):
            yield response.follow(url = self.book + f"{9985615+i}.html", callback=self.parse_chapter, priority = self.priority)
            self.priority -= 1
    
    def parse_chapter(self,response): #on chapter page
        title = response.css('h1.readTitle ::text').get()
        self.doc.add_heading(title)
        content = response.css('div.panel-body[id=htmlContent]::text').getall()
        for para in content:
            p = para.strip() #removing leading and trailing whitespaces
            if p=='':
                continue
            self.doc.add_paragraph(p)
            self.doc.add_paragraph('') #adding a blank line before a new para
        self.doc.add_page_break()
        # below part not understood
        # if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
        #     self.doc.save(f'{self.name}--{((self.count-2)//100)*100} -{self.count}.docx')
        #     self.doc = docx.Document()
        # self.count += 1

    def closed(self,reason):
        self.doc.save(f'{self.name}-{self.start}-{self.end}.docx') 
    
        