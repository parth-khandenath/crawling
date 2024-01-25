import scrapy
import docx
import math

class toyManga(scrapy.Spider):
    name = 'toyManga'
    allowed_domains = ['www.toymanga.com']
    def __init__(self, **kwargs):
        self.start_urls = []
        self.start_chapter = 251 #change here
        self.end_chapter = 305 #change here
        self.book_name = "the president's pregnant, ex-wife" #change here
        start_page=math.ceil(self.start_chapter/50)
        end_page=math.ceil(self.end_chapter/50)
        for i in range(start_page,end_page+1): #change below link
            self.start_urls.append(f'https://www.toymanga.com/book/7000.html?page_num={i}&b=The-President%E2%80%99s-Pregnant,-Ex-Wife')
        self.doc = docx.Document()
        self.priority = 1000000
                
    custom_settings={
        'CONCURRENT_REQUESTS':1
    }
    
    def parse(self,response):  #on a page
        # all_anchors=response.css('.list-charts a').getall()
        all_links=response.css('#list-chapter a::attr("href")').getall()
        all_links=all_links[:50]
        # print(all_links)
        pg_no=int(response.url.split('&')[0].split('=')[1])
        offset=50*(pg_no-1)
        base_url='https://www.toymanga.com'
        for idx in range(1,51):
            chpt_no=offset+idx
            if chpt_no>=self.start_chapter and chpt_no<=self.end_chapter:
                self.priority-=1
                yield response.follow(url=base_url+all_links[idx-1], callback=self.parse_chapter,priority=self.priority)
            elif chpt_no>self.end_chapter:
                break

    def parse_chapter(self,response):
        chp_title=response.css('.c_title::text').get()
        chpt_paras=response.css('#chapter-content div::text').getall()
        rest_paras=response.css('#chapter-content p::text').getall()
        chpt_paras.extend(rest_paras)

        self.doc.add_heading(chp_title)
        for para in chpt_paras:
            para=para.strip()
            if para=='':
                continue
            self.doc.add_paragraph(para)
            self.doc.add_paragraph('')
        self.doc.add_page_break()

    def closed(self,reason):
        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')
    