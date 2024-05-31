import scrapy
import os
from docx import Document
import docx

class SevenHaojidi(scrapy.Spider):
    name = "SevenHaojidi"
    allowed_domains = ["7haojidi.com"]
    start_urls = ["https://7haojidi.com/book/42837/12565204.html"]  #change here    #put url of start chapter
    custom_settings = {
        'ROBOTSTXT_OBEY': False,
        'CONCURRENT_REQUESTS':1,
    }
    def __init__(self):
        self.doc = docx.Document()
        self.start_chp = 1  #change here
        self.end_chp = 228   #change here
        self.curr_chp = self.start_chp
        self.start = self.start_chp
        self.bookname = "Sold AS The alpha King's Breeder new"  #change here
        self.doc = docx.Document()
        self.priority = 100000000


    def parse(self, response):   #on chapters_page
        heading=response.css(".h1.title::text").get()
        paras = response.css('#content p').getall()
        
        #Add data to file
        self.doc.add_heading(heading, level=1)
        for p in paras:
            p=p.strip()
            if p=='':
                continue
            self.doc.add_paragraph(p)
        self.doc.add_page_break()
        
        if ((self.count) % 100 == 0 and (self.count - self.start) != 0) or self.count == self.end:
                    self.doc.save(f'{self.bookname}-{((self.count-2)//100)*100 +1}-{self.count}.docx')
                    self.doc = docx.Document()

        self.count += 1

        if self.count <= self.end:
            relative_url= response.css('.decoration-solid ::attr(href)').getall()[1]
            if self.count<50:
                next_chapter_url= "https://www.novelsquare.com" + relative_url
            else:
                next_chapter_url=   relative_url
            print("NEXT CHAPTER AT ", next_chapter_url)
            yield response.follow(next_chapter_url, callback=self.parse)
        else:
                print("ELSE MAI HAI HUM")
                self.doc.save(f'{self.bookname}-{self.count}.docx')
                self.error_end = True
    
    def closed(self,reason):
        print(self.count)

        print(reason)
        print("CLOSEDDDDDDDDDDDDDDDD")
        self.doc.save('trial.docx') 