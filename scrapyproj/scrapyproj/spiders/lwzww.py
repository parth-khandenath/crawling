import scrapy
import docx
import time
import re
import requests
import time
import undetected_chromedriver as uc
from selenium.common.exceptions import WebDriverException
from selenium.webdriver.common.by import By
# chapters 88-91    189-192
class lwzww(scrapy.Spider):
    name = 'lwzww'
    allowed_domains = ['lwzww.com']

    def __init__(self, **kwargs):
        self.start_urls = []
        self.start_chapter = 1 #change here
        self.end_chapter = 20 #change here
        self.book_name = "Death heartbeat" #change here
        self.start_page= (self.start_chapter-1)//100 + 1
        self.end_page= (self.end_chapter-1)//100 + 1
        for i in range(self.start_page,self.end_page+1): #change below link
            self.start_urls.append(f'http://www.lwzww.com/book/5ilf/{str(i)}/')
        self.doc = docx.Document()

        headers = {
            'accept': 'application/json, text/plain, */*',
            'accept-language': 'en-GB,en-US;q=0.9,en;q=0.8',
            'content-type': 'application/json',
            'cookie': 'qq_domain_video_guid_verify=7f3ccf505389bdc2; _qimei_uuid42=183010e113a1002f9862363182dec709edd78e889d; _qimei_fingerprint=1c609ff4531c6a835b185cfa96d323f7; pgv_pvid=2097815516; video_platform=2; video_guid=7f3ccf505389bdc2; _qimei_q36=; _qimei_h38=2db878989862363182dec70902000001218301; pgv_info=ssid=s3417771617',
            'origin': 'http://www.lwzww.com',
            'referer': 'http://www.lwzww.com/',
            'sec-ch-ua': '"Chromium";v="122", "Not(A:Brand";v="24", "Google Chrome";v="122"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"Windows"',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-site',
            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
        }

        options = uc.ChromeOptions()
        options.add_argument(f"--header={headers}")
        self.driver = uc.Chrome(options=options)
                
    custom_settings={
        'CONCURRENT_REQUESTS':1,
        'headers': {
            'Referer': 'http://www.lwzww.com',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
            }
    }
    
    def parse(self,response): 
        a_tags=response.css('.chapter-5c993fac a::attr(onclick)').getall()
        chpt_links=[]
        for a in a_tags:
            match = re.search(r"location\.href='(.*?)'", a)
            if match:
                url = match.group(1)
                chpt_links.append('http://www.lwzww.com'+url)
        for l in chpt_links:
            yield scrapy.Request(url=l,callback=self.parse_chapter)

    def parse_chapter(self,response):
        title=response.css('.h1-5c993fac::text').get()
        paras=response.css('#chaptercontent p::text').getall()
        self.doc.add_heading(title)
        for para in paras:
            para=para.strip()
            if para=='':
                continue
            self.doc.add_paragraph(para)
            self.doc.add_paragraph('')

    def closed(self,reason):
        self.doc.save(f'{self.book_name}-{self.start_chapter}-{self.end_chapter}.docx')
    
# https://hotnovelpub.com/hotnovelpub-contract-marriage-the-replacement-groom/chapter-1.2669018